from langchain_core.messages import AIMessage
from langchain_core.tools import tool
from docx import Document
from langgraph.prebuilt import ToolNode
import os
from llm_client.llm_client import LLMProxyChatOpenAI
from memory.memory import messages
import chainlit as cl
import json




def load_market_analysis_prompt():
    prompt_path = os.path.join(os.path.dirname(__file__), "../prompts", "market_analysis.txt")
    try:
        with open(prompt_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except FileNotFoundError:
        return "You are a helpful AI assistant."
    
    
llm = LLMProxyChatOpenAI()
market_analysis_prompt = load_market_analysis_prompt()

@tool
def generate_market_report() -> str:
    """Generate market report when its explicitly requested"""
    messages.append({"role": "system", "content": market_analysis_prompt})
    messages.extend(cl.chat_context.to_openai())
    print(messages)
    response =llm.generate([messages])
    filename="market_report.docx"
    paragraph = response.generations[0][0].text
    if os.path.exists(filename):
        doc=Document(filename)
    else:
        doc = Document()
    lines = paragraph.split('\n')
    
    for line in lines:
        line = line.strip()
        # Check for headers and format accordingly
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("**") and line.endswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(line[2:-2])
            run.bold = True
        elif line.find("**") != -1:
            para = doc.add_paragraph()
            add_bold_text(para, line)                                                                   
        elif line.startswith("- "):
            doc.add_paragraph(line[2:])
        elif line.startswith("1. "):
            doc.add_paragraph(line[3:])
        elif line:
            doc.add_paragraph(line)
            
    doc.save(filename)
    return " "
        
    
    
def add_bold_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            paragraph.add_run(part)
        else:
            paragraph.add_run(part).bold = True
        
# création des tools 

folder_path = "data"  # Remplacez par le chemin de votre dossier contenant les fichiers JSON
key_to_extract_sector = "créations_entreprises_par_secteur"
key_to_extract_statut = "création_entreprises_par_nature_juridique"
key_to_extract_implantation = "création_entreprises_par_implantation" 
key_to_extract_region = "création_entreprises_par_région" 

@tool
def fetch_sector_data(folder_path, key_to_extract_sector):
    """ Récupérer les données d'entreprise par secteur """
    result = {}
    
    # Vérification de l'existence du dossier
    if not os.path.exists(folder_path):
        return {"error": "Le dossier spécifié n'existe pas."}

    # Parcourt tous les fichiers du dossier
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".json"):  # Limité aux fichiers JSON
            file_path = os.path.join(folder_path, file_name)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                    # Recherche et extraction des données associées à la clé
                    if key_to_extract_sector in data:
                        secteur_data = data[key_to_extract_sector]
                        extracted_data = []

                        # Parcourt la liste pour récupérer les secteurs et les nombres
                        for entry in secteur_data:
                            if isinstance(entry, dict):
                                secteur = entry.get("secteur", "Inconnu")
                                nombre_de_creations = entry.get("nombre_de_creations", "Inconnu")
                                extracted_data.append({
                                    "secteur": secteur,
                                    "nombre_de_creations": nombre_de_creations
                                })

                        # Ajoute les données extraites au résultat
                        if extracted_data:
                            result[file_name] = extracted_data

            except (json.JSONDecodeError, FileNotFoundError) as e:
                result[file_name] = f"Erreur lors de la lecture du fichier : {e}"
    
    return print(json.dumps(result, ensure_ascii=False, indent=4))

@tool
def fetch_legal_status_data(folder_path, key_to_extract_statut):
    """
    Récupperer les données d'entreprise par statut juridique
    """
    result = {}
    
    # Vérification de l'existence du dossier
    if not os.path.exists(folder_path):
        return {"error": "Le dossier spécifié n'existe pas."}

    # Parcourt tous les fichiers du dossier
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".json"):  # Limité aux fichiers JSON
            file_path = os.path.join(folder_path, file_name)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                    # Recherche et extraction des données associées à la clé
                    if key_to_extract_statut in data:
                        statut_data = data[key_to_extract_statut]
                        extracted_data = []

                        # Parcourt la liste pour récupérer les statuts juridiques et les nombres de créations
                        for entry in statut_data:
                            if isinstance(entry, dict):
                                statut_juridique = entry.get("statut_juridique", "Inconnu")
                                nombre_de_creations = entry.get("nombre_de_creations", "Inconnu")
                                extracted_data.append({
                                    "statut_juridique": statut_juridique,
                                    "nombre_de_creations": nombre_de_creations
                                })

                        # Ajoute les données extraites au résultat
                        if extracted_data:
                            result[file_name] = extracted_data

            except (json.JSONDecodeError, FileNotFoundError) as e:
                result[file_name] = f"Erreur lors de la lecture du fichier : {e}"
    
    return print(json.dumps(result, ensure_ascii=False, indent=4))


@tool
def fetch_implatation_data(key_to_extract_implantation):
    """
    Récupperer les données d'entreprise par implantation
    """
    result = {}

    # Vérification de l'existence du dossier
    if not os.path.exists(folder_path):
        return {"error": "Le dossier spécifié n'existe pas."}

    # Parcourt tous les fichiers du dossier
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".json"):  # Limité aux fichiers JSON
            file_path = os.path.join(folder_path, file_name)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                    # Recherche et extraction des données associées à la clé
                    if key_to_extract_implantation in data:
                        implantation_data = data[key_to_extract_implantation]
                        extracted_data = []

                        # Parcourt la liste pour récupérer les informations sur les implantations
                        for entry in implantation_data:
                            if isinstance(entry, dict):
                                implantation = entry.get("implantation", "Inconnu")
                                nombre_de_creations = entry.get("nombre_de_creations", "Inconnu")
                                extracted_data.append({
                                    "implantation": implantation,
                                    "nombre_de_creations": nombre_de_creations
                                })

                        # Ajoute les données extraites au résultat
                        if extracted_data:
                            result[file_name] = extracted_data

            except (json.JSONDecodeError, FileNotFoundError) as e:
                result[file_name] = f"Erreur lors de la lecture du fichier : {e}"

    return print(json.dumps(result, ensure_ascii=False, indent=4))


@tool
def fetch_region_data():
    """
    récupérer les données d'entreprise par région
    """
    result = {}
    
    # Vérification de l'existence du dossier
    if not os.path.exists(folder_path):
        return {"error": "Le dossier spécifié n'existe pas."}

    # Parcourt tous les fichiers du dossier
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".json"):  # Limité aux fichiers JSON
            file_path = os.path.join(folder_path, file_name)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                    # Recherche et extraction des données associées à la clé
                    if key_to_extract_region in data:
                        region_data = data[key_to_extract_region]
                        extracted_data = []

                        # Parcourt la liste pour récupérer les informations sur les régions
                        for entry in region_data:
                            if isinstance(entry, dict):
                                region = entry.get("région", "Inconnu")
                                nombre_de_creations = entry.get("nombre_de_creations", "Inconnu")
                                extracted_data.append({
                                    "région": region,
                                    "nombre_de_creations": nombre_de_creations
                                })

                        # Ajoute les données extraites au résultat
                        if extracted_data:
                            result[file_name] = extracted_data

            except (json.JSONDecodeError, FileNotFoundError) as e:
                result[file_name] = f"Erreur lors de la lecture du fichier : {e}"
    
    return print(json.dumps(result, ensure_ascii=False, indent=4))


tools = [generate_market_report, fetch_sector_data, fetch_legal_status_data, fetch_implatation_data, fetch_region_data]
tool_node = ToolNode(tools)